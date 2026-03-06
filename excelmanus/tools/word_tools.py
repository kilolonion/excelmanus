# -*- coding: utf-8 -*-
"""Word 文档工具：提供 .docx 读取、写入、检查和搜索能力。"""

from __future__ import annotations

import json
import os
import re
from pathlib import Path
from typing import Any

from excelmanus.logger import get_logger
from excelmanus.tools._guard_ctx import get_guard as _get_ctx_guard
from excelmanus.tools._helpers import check_file_exists
from excelmanus.tools.registry import ToolDef

logger = get_logger("tools.word")

_WORD_SUFFIXES: frozenset[str] = frozenset({".docx"})

# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------


def _resolve_path(file_path: str) -> tuple[Path, str | None]:
    """解析并校验文件路径，返回 (safe_path, error_json)。"""
    guard = _get_ctx_guard()
    if guard is None:
        return Path(file_path), None
    safe = guard.resolve_and_validate(file_path)
    err = check_file_exists(safe, file_path, guard)
    return safe, err


def _ensure_docx(file_path: str) -> str | None:
    """如果文件不是 .docx 返回错误 JSON，否则返回 None。"""
    if not file_path.lower().endswith(".docx"):
        return json.dumps({
            "error": "仅支持 .docx 格式文件，当前文件不是 .docx",
            "file_path": file_path,
        }, ensure_ascii=False)
    return None


def _open_docx(safe_path: Path):  # -> docx.Document
    from docx import Document
    return Document(str(safe_path))


def _para_style_name(para) -> str:
    """获取段落样式名，回退到 'Normal'。"""
    try:
        return para.style.name or "Normal"
    except Exception:
        return "Normal"


def _heading_level(para) -> int | None:
    """如果段落是 Heading 样式，返回其级别 (1-9)，否则返回 None。"""
    style = _para_style_name(para)
    if style.startswith("Heading"):
        try:
            return int(style.split()[-1])
        except (ValueError, IndexError):
            pass
    if style == "Title":
        return 0
    return None


def _run_to_dict(run) -> dict[str, Any]:
    """将 Run 对象转为精简字典。"""
    d: dict[str, Any] = {"text": run.text}
    fmt: dict[str, Any] = {}
    if run.bold:
        fmt["bold"] = True
    if run.italic:
        fmt["italic"] = True
    if run.underline:
        fmt["underline"] = True
    if run.font.size:
        fmt["size_pt"] = round(run.font.size.pt, 1)
    if run.font.name:
        fmt["font"] = run.font.name
    if run.font.color and run.font.color.rgb:
        fmt["color"] = str(run.font.color.rgb)
    if fmt:
        d["format"] = fmt
    return d


def _table_to_dict(table) -> dict[str, Any]:
    """将 Table 对象转为字典。"""
    rows_data: list[list[str]] = []
    for row in table.rows:
        rows_data.append([cell.text.strip() for cell in row.cells])
    return {
        "rows": len(table.rows),
        "columns": len(table.columns),
        "data": rows_data,
    }


# ---------------------------------------------------------------------------
# read_word
# ---------------------------------------------------------------------------


def read_word(
    file_path: str,
    *,
    offset: int = 0,
    max_paragraphs: int = 100,
    include_format: bool = False,
    include_tables: bool = True,
) -> str:
    """读取 Word 文档段落和表格。

    Args:
        file_path: .docx 文件路径
        offset: 起始段落索引（0-based）
        max_paragraphs: 最多返回段落数
        include_format: 是否包含行内格式信息（加粗/斜体/字号等）
        include_tables: 是否包含表格数据
    """
    fmt_err = _ensure_docx(file_path)
    if fmt_err:
        return fmt_err

    safe_path, err = _resolve_path(file_path)
    if err:
        return err

    try:
        doc = _open_docx(safe_path)
    except Exception as exc:
        return json.dumps({"error": f"无法打开文档: {exc}"}, ensure_ascii=False)

    total = len(doc.paragraphs)
    end = min(offset + max_paragraphs, total)
    paragraphs: list[dict[str, Any]] = []

    for i in range(offset, end):
        para = doc.paragraphs[i]
        entry: dict[str, Any] = {
            "index": i,
            "text": para.text,
            "style": _para_style_name(para),
        }
        level = _heading_level(para)
        if level is not None:
            entry["heading_level"] = level

        if include_format and para.runs:
            entry["runs"] = [_run_to_dict(r) for r in para.runs if r.text]

        paragraphs.append(entry)

    result: dict[str, Any] = {
        "file": file_path,
        "total_paragraphs": total,
        "offset": offset,
        "returned": len(paragraphs),
        "truncated": end < total,
        "paragraphs": paragraphs,
    }

    if include_tables:
        tables = []
        for idx, tbl in enumerate(doc.tables):
            t = _table_to_dict(tbl)
            t["index"] = idx
            tables.append(t)
        result["tables"] = tables
        result["total_tables"] = len(doc.tables)

    return json.dumps(result, ensure_ascii=False, default=str)


# ---------------------------------------------------------------------------
# write_word
# ---------------------------------------------------------------------------


def write_word(
    file_path: str,
    *,
    operations: list[dict[str, Any]],
) -> str:
    """对 Word 文档执行写入操作。

    operations 列表中每个操作支持:
      - {"action": "replace", "index": 段落索引, "text": "新内容"}
      - {"action": "insert_after", "index": 段落索引, "text": "新内容", "style": "Normal"}
      - {"action": "append", "text": "新内容", "style": "Normal"}
      - {"action": "delete", "index": 段落索引}
    """
    fmt_err = _ensure_docx(file_path)
    if fmt_err:
        return fmt_err

    safe_path, err = _resolve_path(file_path)
    if err:
        return err

    try:
        doc = _open_docx(safe_path)
    except Exception as exc:
        return json.dumps({"error": f"无法打开文档: {exc}"}, ensure_ascii=False)

    applied: list[str] = []
    errors: list[str] = []

    for op in operations:
        action = op.get("action", "")
        idx = op.get("index")
        text = op.get("text", "")
        style = op.get("style")

        try:
            if action == "replace":
                if idx is None or idx < 0 or idx >= len(doc.paragraphs):
                    errors.append(f"replace: 段落索引 {idx} 超出范围 (0-{len(doc.paragraphs)-1})")
                    continue
                para = doc.paragraphs[idx]
                para.clear()
                para.add_run(text)
                if style:
                    para.style = doc.styles[style]
                applied.append(f"replace paragraph {idx}")

            elif action == "insert_after":
                if idx is None or idx < 0 or idx >= len(doc.paragraphs):
                    errors.append(f"insert_after: 段落索引 {idx} 超出范围")
                    continue
                ref_para = doc.paragraphs[idx]
                from docx.oxml.ns import qn
                new_p = ref_para._element.makeelement(qn("w:p"), {})
                ref_para._element.addnext(new_p)
                from docx.text.paragraph import Paragraph
                new_para = Paragraph(new_p, ref_para._parent)
                new_para.add_run(text)
                if style:
                    new_para.style = doc.styles[style]
                applied.append(f"insert_after paragraph {idx}")

            elif action == "append":
                para = doc.add_paragraph(text, style=style)
                applied.append("append paragraph")

            elif action == "delete":
                if idx is None or idx < 0 or idx >= len(doc.paragraphs):
                    errors.append(f"delete: 段落索引 {idx} 超出范围")
                    continue
                p_element = doc.paragraphs[idx]._element
                p_element.getparent().remove(p_element)
                applied.append(f"delete paragraph {idx}")

            else:
                errors.append(f"未知操作: {action}")
        except Exception as exc:
            errors.append(f"{action} index={idx}: {exc}")

    try:
        doc.save(str(safe_path))
    except Exception as exc:
        return json.dumps({"error": f"保存文档失败: {exc}"}, ensure_ascii=False)

    result: dict[str, Any] = {
        "file": file_path,
        "applied": applied,
        "applied_count": len(applied),
    }
    if errors:
        result["errors"] = errors
    return json.dumps(result, ensure_ascii=False)


# ---------------------------------------------------------------------------
# inspect_word
# ---------------------------------------------------------------------------


def inspect_word(
    file_path: str | None = None,
    *,
    file_paths: list[str] | None = None,
    directory: str = ".",
) -> str:
    """检查 Word 文档的结构概览。

    可传单个 file_path 或多个 file_paths；若都为空则扫描 directory。
    """
    paths: list[str] = []
    if file_paths:
        paths = file_paths
    elif file_path:
        paths = [file_path]
    else:
        guard = _get_ctx_guard()
        if guard:
            root = guard.resolve_and_validate(directory)
        else:
            root = Path(directory)
        if root.is_dir():
            for f in sorted(root.iterdir()):
                if f.is_file() and f.suffix.lower() in _WORD_SUFFIXES:
                    try:
                        rel = str(f.relative_to(guard.workspace_root)) if guard else str(f)
                        paths.append(rel)
                    except ValueError:
                        paths.append(f.name)
        if not paths:
            return json.dumps({"error": f"目录 '{directory}' 下未找到 Word 文件"}, ensure_ascii=False)

    results: list[dict[str, Any]] = []

    for fp in paths[:10]:
        if not fp.lower().endswith(".docx"):
            results.append({"file": fp, "error": "仅支持 .docx 格式"})
            continue

        safe_path, err = _resolve_path(fp)
        if err:
            results.append({"file": fp, "error": json.loads(err).get("error", err)})
            continue

        try:
            doc = _open_docx(safe_path)
        except Exception as exc:
            results.append({"file": fp, "error": str(exc)})
            continue

        headings: list[dict[str, Any]] = []
        for i, para in enumerate(doc.paragraphs):
            level = _heading_level(para)
            if level is not None:
                headings.append({"index": i, "level": level, "text": para.text.strip()})

        sections = []
        for idx, sec in enumerate(doc.sections):
            sections.append({
                "index": idx,
                "width_cm": round(sec.page_width.cm, 1) if sec.page_width else None,
                "height_cm": round(sec.page_height.cm, 1) if sec.page_height else None,
                "orientation": "landscape" if sec.orientation else "portrait",
            })

        info: dict[str, Any] = {
            "file": fp,
            "total_paragraphs": len(doc.paragraphs),
            "total_tables": len(doc.tables),
            "total_sections": len(doc.sections),
            "headings": headings,
            "sections": sections,
            "size_bytes": safe_path.stat().st_size,
        }

        core = doc.core_properties
        if core.title:
            info["title"] = core.title
        if core.author:
            info["author"] = core.author

        results.append(info)

    if len(results) == 1:
        return json.dumps(results[0], ensure_ascii=False, default=str)
    return json.dumps({"files": results}, ensure_ascii=False, default=str)


# ---------------------------------------------------------------------------
# search_word
# ---------------------------------------------------------------------------


def search_word(
    query: str,
    *,
    file_path: str | None = None,
    file_paths: list[str] | None = None,
    match_mode: str = "contains",
    case_sensitive: bool = False,
    max_results: int = 50,
) -> str:
    """在 Word 文档中搜索文本。

    match_mode: contains | exact | regex | startswith
    """
    paths: list[str] = []
    if file_paths:
        paths = file_paths
    elif file_path:
        paths = [file_path]
    else:
        return json.dumps({"error": "必须提供 file_path 或 file_paths"}, ensure_ascii=False)

    flags = 0 if case_sensitive else re.IGNORECASE
    matches: list[dict[str, Any]] = []
    errors: list[dict[str, Any]] = []
    inspected_files = 0

    for fp in paths[:10]:
        fmt_err = _ensure_docx(fp)
        if fmt_err:
            errors.append(json.loads(fmt_err))
            continue

        safe_path, err = _resolve_path(fp)
        if err:
            errors.append(json.loads(err))
            continue

        try:
            doc = _open_docx(safe_path)
        except Exception as exc:
            errors.append({
                "error": f"无法打开文档: {exc}",
                "file_path": fp,
            })
            continue

        inspected_files += 1

        for i, para in enumerate(doc.paragraphs):
            text = para.text
            hit = False

            if match_mode == "exact":
                hit = (text == query) if case_sensitive else (text.lower() == query.lower())
            elif match_mode == "startswith":
                hit = text.startswith(query) if case_sensitive else text.lower().startswith(query.lower())
            elif match_mode == "regex":
                try:
                    hit = bool(re.search(query, text, flags))
                except re.error:
                    hit = False
            else:  # contains
                hit = (query in text) if case_sensitive else (query.lower() in text.lower())

            if hit:
                matches.append({
                    "file": fp,
                    "paragraph_index": i,
                    "style": _para_style_name(para),
                    "text": text[:500],
                })
                if len(matches) >= max_results:
                    break

        if len(matches) >= max_results:
            break

    if len(paths[:10]) == 1 and inspected_files == 0 and errors:
        return json.dumps(errors[0], ensure_ascii=False, default=str)

    result: dict[str, Any] = {
        "query": query,
        "match_mode": match_mode,
        "total_matches": len(matches),
        "matches": matches,
    }
    if errors:
        result["errors"] = errors
    return json.dumps(result, ensure_ascii=False, default=str)


# ---------------------------------------------------------------------------
# get_tools
# ---------------------------------------------------------------------------


def get_tools() -> list[ToolDef]:
    return [
        ToolDef(
            name="read_word",
            description=(
                "读取 Word (.docx) 文档的段落内容和表格。"
                "返回段落文本、样式名称和标题层级，可选包含行内格式（加粗/斜体/字号等）。"
                "支持分页：通过 offset 和 max_paragraphs 控制读取范围。"
            ),
            input_schema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Word 文件路径（相对于工作目录）",
                    },
                    "offset": {
                        "type": "integer",
                        "description": "起始段落索引（0-based），默认 0",
                        "default": 0,
                        "minimum": 0,
                    },
                    "max_paragraphs": {
                        "type": "integer",
                        "description": "最多返回段落数，默认 100",
                        "default": 100,
                        "minimum": 1,
                    },
                    "include_format": {
                        "type": "boolean",
                        "description": "是否包含行内格式信息（加粗/斜体/字号等），默认 false",
                        "default": False,
                    },
                    "include_tables": {
                        "type": "boolean",
                        "description": "是否包含表格数据，默认 true",
                        "default": True,
                    },
                },
                "required": ["file_path"],
                "additionalProperties": False,
            },
            func=read_word,
            max_result_chars=8000,
            write_effect="none",
        ),
        ToolDef(
            name="write_word",
            description=(
                "对 Word (.docx) 文档执行写入操作。"
                "支持的操作：replace（替换段落内容）、insert_after（在段落后插入）、"
                "append（追加段落）、delete（删除段落）。"
                "每次调用可包含多个操作，按顺序执行。"
            ),
            input_schema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Word 文件路径（相对于工作目录）",
                    },
                    "operations": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "action": {
                                    "type": "string",
                                    "enum": ["replace", "insert_after", "append", "delete"],
                                },
                                "index": {
                                    "type": "integer",
                                    "description": "目标段落索引（replace/insert_after/delete 需要）",
                                },
                                "text": {
                                    "type": "string",
                                    "description": "新内容（replace/insert_after/append 需要）",
                                },
                                "style": {
                                    "type": "string",
                                    "description": "段落样式名（可选，如 Heading 1, Normal）",
                                },
                            },
                            "required": ["action"],
                        },
                        "description": "写入操作列表",
                    },
                },
                "required": ["file_path", "operations"],
                "additionalProperties": False,
            },
            func=write_word,
            write_effect="workspace_write",
        ),
        ToolDef(
            name="inspect_word",
            description=(
                "检查 Word (.docx) 文档的结构：标题树、段落数、表格数、节数、页面设置。"
                "可传单个 file_path、多个 file_paths，或扫描 directory 下所有 .docx 文件。"
                "适用于了解文档结构后再进行精准编辑。"
            ),
            input_schema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "单个 Word 文件路径",
                    },
                    "file_paths": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "多个文件路径（最多 10 个）",
                    },
                    "directory": {
                        "type": "string",
                        "description": "扫描目录（file_path 和 file_paths 均为空时使用），默认当前目录",
                        "default": ".",
                    },
                },
                "required": [],
                "additionalProperties": False,
            },
            func=inspect_word,
            max_result_chars=6000,
            write_effect="none",
        ),
        ToolDef(
            name="search_word",
            description=(
                "在 Word (.docx) 文档中搜索文本，返回匹配的段落及位置。"
                "支持多种匹配模式：contains（包含）、exact（精确）、regex（正则）、startswith（前缀）。"
                "可同时搜索多个文件。"
            ),
            input_schema={
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "搜索字符串或正则表达式",
                    },
                    "file_path": {
                        "type": "string",
                        "description": "单个 Word 文件路径",
                    },
                    "file_paths": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "多个文件路径（最多 10 个）",
                    },
                    "match_mode": {
                        "type": "string",
                        "enum": ["contains", "exact", "regex", "startswith"],
                        "description": "匹配模式，默认 contains",
                        "default": "contains",
                    },
                    "case_sensitive": {
                        "type": "boolean",
                        "description": "是否区分大小写，默认 false",
                        "default": False,
                    },
                    "max_results": {
                        "type": "integer",
                        "description": "最大返回匹配数，默认 50",
                        "default": 50,
                        "minimum": 1,
                    },
                },
                "required": ["query"],
                "additionalProperties": False,
            },
            func=search_word,
            max_result_chars=8000,
            write_effect="none",
        ),
    ]
