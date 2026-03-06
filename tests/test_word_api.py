"""Word file API tests."""

from __future__ import annotations

from contextlib import contextmanager
from pathlib import Path

import pytest
from docx import Document
from httpx import AsyncClient

from tests.test_api import _make_transport, _setup_api_globals, _test_config


pytestmark = pytest.mark.asyncio


def _make_test_doc(path: Path, paragraphs=None) -> Path:
    doc = Document()
    for text in (paragraphs or ["Hello", "World"]):
        doc.add_paragraph(text)
    doc.save(path)
    return path


@contextmanager
def _api_transport(tmp_path: Path):
    config = _test_config(workspace_root=str(tmp_path))
    with _setup_api_globals(config=config):
        yield _make_transport()


class TestWordSnapshot:
    async def test_snapshot_returns_document_json(self, tmp_path: Path) -> None:
        _make_test_doc(tmp_path / "report.docx", ["Intro", "Summary"])

        with _api_transport(tmp_path) as transport:
            async with AsyncClient(transport=transport, base_url="http://test") as client:
                response = await client.get("/api/v1/files/word/snapshot", params={"path": "report.docx"})

        data = response.json()
        assert response.status_code == 200
        assert data["file"] == "report.docx"
        assert data["total_paragraphs"] == 2
        assert data["paragraphs"][0]["text"] == "Intro"

    async def test_snapshot_handles_empty_document(self, tmp_path: Path) -> None:
        Document().save(tmp_path / "empty.docx")

        with _api_transport(tmp_path) as transport:
            async with AsyncClient(transport=transport, base_url="http://test") as client:
                response = await client.get("/api/v1/files/word/snapshot", params={"path": "empty.docx"})

        data = response.json()
        assert response.status_code == 200
        assert data["total_paragraphs"] == 0
        assert data["paragraphs"] == []

    async def test_snapshot_missing_file_returns_error(self, tmp_path: Path) -> None:
        with _api_transport(tmp_path) as transport:
            async with AsyncClient(transport=transport, base_url="http://test") as client:
                response = await client.get("/api/v1/files/word/snapshot", params={"path": "missing.docx"})

        data = response.json()
        assert response.status_code == 404
        assert "Word" in data["error"]
        assert "不存在" in data["error"]

    async def test_snapshot_rejects_doc_file(self, tmp_path: Path) -> None:
        (tmp_path / "legacy.doc").write_bytes(b"legacy-doc")

        with _api_transport(tmp_path) as transport:
            async with AsyncClient(transport=transport, base_url="http://test") as client:
                response = await client.get("/api/v1/files/word/snapshot", params={"path": "legacy.doc"})

        data = response.json()
        assert response.status_code == 404
        assert "docx" in data["error"].lower()


class TestWordWrite:
    async def test_replace_writes_content(self, tmp_path: Path) -> None:
        path = _make_test_doc(tmp_path / "report.docx")

        with _api_transport(tmp_path) as transport:
            async with AsyncClient(transport=transport, base_url="http://test") as client:
                response = await client.post(
                    "/api/v1/files/word/write",
                    json={"path": "report.docx", "operations": [{"action": "replace", "index": 0, "text": "Updated"}]},
                )

        assert response.status_code == 200
        assert response.json()["applied"] == ["replace paragraph 0"]
        assert [paragraph.text for paragraph in Document(path).paragraphs] == ["Updated", "World"]

    async def test_insert_after_regression(self, tmp_path: Path) -> None:
        path = _make_test_doc(tmp_path / "insert.docx")

        with _api_transport(tmp_path) as transport:
            async with AsyncClient(transport=transport, base_url="http://test") as client:
                response = await client.post(
                    "/api/v1/files/word/write",
                    json={"path": "insert.docx", "operations": [{"action": "insert_after", "index": 0, "text": "Inserted"}]},
                )

        assert response.status_code == 200
        assert response.json()["applied_count"] == 1
        assert [paragraph.text for paragraph in Document(path).paragraphs] == ["Hello", "Inserted", "World"]

    async def test_out_of_range_returns_error(self, tmp_path: Path) -> None:
        path = _make_test_doc(tmp_path / "bounds.docx")

        with _api_transport(tmp_path) as transport:
            async with AsyncClient(transport=transport, base_url="http://test") as client:
                response = await client.post(
                    "/api/v1/files/word/write",
                    json={"path": "bounds.docx", "operations": [{"action": "replace", "index": 9, "text": "Never"}]},
                )

        data = response.json()
        assert response.status_code == 200
        assert "errors" in data
        assert "超出范围" in data["errors"][0]
        assert [paragraph.text for paragraph in Document(path).paragraphs] == ["Hello", "World"]

    async def test_write_rejects_doc_file(self, tmp_path: Path) -> None:
        (tmp_path / "legacy.doc").write_bytes(b"legacy-doc")

        with _api_transport(tmp_path) as transport:
            async with AsyncClient(transport=transport, base_url="http://test") as client:
                response = await client.post(
                    "/api/v1/files/word/write",
                    json={"path": "legacy.doc", "operations": [{"action": "append", "text": "Never"}]},
                )

        data = response.json()
        assert response.status_code == 404
        assert "docx" in data["error"].lower()


class TestWordFile:
    async def test_file_download_rejects_doc_file(self, tmp_path: Path) -> None:
        (tmp_path / "legacy.doc").write_bytes(b"legacy-doc")

        with _api_transport(tmp_path) as transport:
            async with AsyncClient(transport=transport, base_url="http://test") as client:
                response = await client.get("/api/v1/files/word", params={"path": "legacy.doc"})

        data = response.json()
        assert response.status_code == 404
        assert "docx" in data["error"].lower()
