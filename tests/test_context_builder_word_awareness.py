"""ContextBuilder 的 Word 感知测试。"""

from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

from docx import Document
from openpyxl import Workbook

from excelmanus.engine_core.context_builder import ContextBuilder


def _make_builder(tmp_path: Path, *, state: SimpleNamespace | None = None) -> ContextBuilder:
    """构造最小可用的 ContextBuilder。"""
    if state is None:
        state = SimpleNamespace(
            explorer_reports=[],
            _explore_in_progress=False,
            write_operations_log=[],
            finish_task_warned=False,
        )

    engine = SimpleNamespace(
        _state=state,
        state=state,
        config=SimpleNamespace(workspace_root=str(tmp_path)),
        _workspace_root=str(tmp_path),
    )
    builder = ContextBuilder.__new__(ContextBuilder)
    builder._engine = engine
    return builder


class TestPostWriteVerificationHint:
    def test_word_write_uses_read_word_hint(self, tmp_path: Path) -> None:
        state = SimpleNamespace(
            write_operations_log=[
                {"tool_name": "write_word", "file_path": "report.docx", "summary": "replace paragraph 1"},
                {"tool_name": "write_word", "file_path": "report.docx", "summary": "insert paragraph 2"},
                {"tool_name": "write_word", "file_path": "report.docx", "summary": "append summary"},
            ],
            finish_task_warned=False,
        )
        builder = _make_builder(tmp_path, state=state)

        hint = builder._build_post_write_verification_hint()

        assert "read_word" in hint
        assert "read_excel" not in hint

    def test_mixed_write_mentions_excel_and_word_reads(self, tmp_path: Path) -> None:
        state = SimpleNamespace(
            write_operations_log=[
                {"tool_name": "write_to_sheet", "file_path": "sales.xlsx", "sheet": "Sheet1", "summary": "update A1:C10"},
                {"tool_name": "write_word", "file_path": "report.docx", "summary": "replace paragraph 1"},
                {"tool_name": "write_word", "file_path": "report.docx", "summary": "append conclusion"},
            ],
            finish_task_warned=False,
        )
        builder = _make_builder(tmp_path, state=state)

        hint = builder._build_post_write_verification_hint()

        assert "read_word" in hint
        assert "read_excel" in hint or "scan_excel_snapshot" in hint


class TestDiscoverWorkspaceFiles:
    def test_discover_excel_files_on_disk_includes_docx(self, tmp_path: Path) -> None:
        excel_path = tmp_path / "sales.xlsx"
        wb = Workbook()
        wb.save(excel_path)

        word_path = tmp_path / "report.docx"
        doc = Document()
        doc.add_paragraph("hello")
        doc.save(word_path)

        text_path = tmp_path / "notes.txt"
        text_path.write_text("ignore me", encoding="utf-8")

        discovered = ContextBuilder._discover_excel_files_on_disk(str(tmp_path))

        assert str(excel_path) in discovered
        assert str(word_path) in discovered
        assert str(text_path) not in discovered


class TestScanToolHint:
    def test_scan_hint_mentions_inspect_word_when_docx_exists(self, tmp_path: Path) -> None:
        word_path = tmp_path / "report.docx"
        doc = Document()
        doc.add_heading("周报", level=1)
        doc.save(word_path)

        builder = _make_builder(tmp_path)

        original = ContextBuilder._try_auto_prescan
        ContextBuilder._try_auto_prescan = staticmethod(lambda paths, state: False)
        try:
            hint = builder._build_scan_tool_hint()
        finally:
            ContextBuilder._try_auto_prescan = original

        assert "inspect_word" in hint


class TestAutoPrescanWord:
    def test_try_auto_prescan_injects_word_report(self, tmp_path: Path) -> None:
        word_path = tmp_path / "report.docx"
        doc = Document()
        doc.add_heading("项目概览", level=1)
        doc.add_paragraph("第一段")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        doc.save(word_path)

        state = SimpleNamespace(explorer_reports=[])

        result = ContextBuilder._try_auto_prescan([str(word_path)], state)

        assert result is True
        assert len(state.explorer_reports) == 1
        report = state.explorer_reports[0]
        assert report["files"][0]["path"] == str(word_path)
        assert "段落" in report["summary"]
        assert "表格" in report["summary"]
        assert any("项目概览" in finding["detail"] for finding in report["findings"])
