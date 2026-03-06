"""Word router regression tests."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock

from docx import Document
from openpyxl import Workbook

from excelmanus.config import ExcelManusConfig
from excelmanus.skillpacks.router import SkillRouter


def _make_router() -> SkillRouter:
    config = ExcelManusConfig(api_key="test", base_url="http://test", model="test")
    loader = MagicMock()
    loader.get_skillpacks.return_value = {}
    loader.load_all.return_value = {}
    return SkillRouter(config, loader)


class TestWordPathExtraction:
    def test_extracts_docx_path_from_message(self) -> None:
        assert SkillRouter._extract_word_paths("请修改 report.docx 的结论") == ["report.docx"]

    def test_extracts_quoted_docx_path(self) -> None:
        assert SkillRouter._extract_word_paths('请读取 "docs/report.docx"') == ["docs/report.docx"]

    def test_returns_empty_without_docx(self) -> None:
        assert SkillRouter._extract_word_paths("请读取 notes.txt") == []


class TestWriteHintLexical:
    def test_modify_docx_is_may_write(self) -> None:
        assert SkillRouter._classify_write_hint_lexical("修改 report.docx") == "may_write"

    def test_read_excel_message_is_unchanged(self) -> None:
        assert SkillRouter._classify_write_hint_lexical("读取 data.xlsx") == "read_only"


class TestFileStructureContext:
    def test_docx_file_builds_context(self, tmp_path: Path) -> None:
        path = tmp_path / "report.docx"
        doc = Document()
        doc.add_heading("Overview", level=1)
        doc.add_paragraph("Body")
        doc.save(path)

        text, sheet_count, max_rows = _make_router()._build_file_structure_context_sync(
            candidate_file_paths=[str(path)]
        )

        assert "[文件结构预览]" in text
        assert "report.docx" in text
        assert "Overview" in text
        assert sheet_count == 0
        assert max_rows == 0

    def test_xlsx_file_still_builds_context(self, tmp_path: Path) -> None:
        path = tmp_path / "sales.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws.append(["Name", "Value"])
        ws.append(["A", 1])
        wb.save(path)

        text, sheet_count, max_rows = _make_router()._build_file_structure_context_sync(
            candidate_file_paths=[str(path)]
        )

        assert "[文件结构预览]" in text
        assert "sales.xlsx" in text
        assert "Sheet1" in text
        assert sheet_count == 1
        assert max_rows >= 2
