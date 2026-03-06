"""Word tool real-I/O regression tests."""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt, RGBColor

from excelmanus.database import Database
from excelmanus.file_registry import FileRegistry
from excelmanus.security.guard import FileAccessGuard
from excelmanus.tools._guard_ctx import reset_guard, set_guard
from excelmanus.tools.word_tools import (
    _ensure_docx,
    _heading_level,
    _run_to_dict,
    inspect_word,
    read_word,
    search_word,
    write_word,
)


@pytest.fixture(autouse=True)
def _set_guard(tmp_path: Path):
    token = set_guard(FileAccessGuard(str(tmp_path)))
    try:
        yield
    finally:
        reset_guard(token)


def _make_test_doc(path, paragraphs=None):
    doc = Document()
    for text in (paragraphs or ["Hello", "World"]):
        doc.add_paragraph(text)
    doc.save(str(path))
    return path


def _paragraph_texts(path: Path) -> list[str]:
    return [paragraph.text for paragraph in Document(path).paragraphs]


class TestReadWord:
    def test_basic_read(self, tmp_path: Path) -> None:
        path = _make_test_doc(tmp_path / "report.docx", ["Intro", "Summary"])
        doc = Document(path)
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Metric"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Revenue"
        table.cell(1, 1).text = "42"
        doc.save(path)

        result = json.loads(read_word(path.name))

        assert result["file"] == path.name
        assert result["total_paragraphs"] == 2
        assert result["returned"] == 2
        assert result["paragraphs"][0]["text"] == "Intro"
        assert result["total_tables"] == 1
        assert result["tables"][0]["data"][1] == ["Revenue", "42"]

    def test_offset_pagination(self, tmp_path: Path) -> None:
        path = _make_test_doc(tmp_path / "paged.docx", ["A", "B", "C"])

        result = json.loads(read_word(path.name, offset=1, max_paragraphs=1))

        assert result["offset"] == 1
        assert result["returned"] == 1
        assert result["truncated"] is True
        assert result["paragraphs"] == [{"index": 1, "text": "B", "style": "Normal"}]

    def test_include_format(self, tmp_path: Path) -> None:
        path = tmp_path / "format.docx"
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Styled")
        run.bold = True
        run.italic = True
        run.font.size = Pt(12)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x11, 0x22, 0x33)
        doc.save(path)

        result = json.loads(read_word(path.name, include_format=True))

        run_data = result["paragraphs"][0]["runs"][0]
        assert run_data["text"] == "Styled"
        assert run_data["format"]["bold"] is True
        assert run_data["format"]["italic"] is True
        assert run_data["format"]["size_pt"] == 12.0
        assert run_data["format"]["font"] == "Calibri"
        assert run_data["format"]["color"] == "112233"

    def test_include_tables_false_omits_tables(self, tmp_path: Path) -> None:
        path = _make_test_doc(tmp_path / "no_tables.docx")
        doc = Document(path)
        doc.add_table(rows=1, cols=1).cell(0, 0).text = "hidden"
        doc.save(path)

        result = json.loads(read_word(path.name, include_tables=False))

        assert "tables" not in result
        assert "total_tables" not in result

    def test_empty_document(self, tmp_path: Path) -> None:
        path = tmp_path / "empty.docx"
        Document().save(path)

        result = json.loads(read_word(path.name))

        assert result["total_paragraphs"] == 0
        assert result["returned"] == 0
        assert result["paragraphs"] == []
        assert result["truncated"] is False


class TestWriteWord:
    def test_replace(self, tmp_path: Path) -> None:
        path = _make_test_doc(tmp_path / "replace.docx")

        result = json.loads(
            write_word(path.name, operations=[{"action": "replace", "index": 0, "text": "Updated"}])
        )

        assert result["applied"] == ["replace paragraph 0"]
        assert _paragraph_texts(path) == ["Updated", "World"]

    def test_insert_after_regression(self, tmp_path: Path) -> None:
        path = _make_test_doc(tmp_path / "insert.docx")

        result = json.loads(
            write_word(path.name, operations=[{"action": "insert_after", "index": 0, "text": "Inserted"}])
        )

        assert result["applied_count"] == 1
        assert "errors" not in result
        assert _paragraph_texts(path) == ["Hello", "Inserted", "World"]

    def test_append(self, tmp_path: Path) -> None:
        path = _make_test_doc(tmp_path / "append.docx")

        result = json.loads(write_word(path.name, operations=[{"action": "append", "text": "Tail"}]))

        assert result["applied"] == ["append paragraph"]
        assert _paragraph_texts(path) == ["Hello", "World", "Tail"]

    def test_delete(self, tmp_path: Path) -> None:
        path = _make_test_doc(tmp_path / "delete.docx")

        result = json.loads(write_word(path.name, operations=[{"action": "delete", "index": 0}]))

        assert result["applied"] == ["delete paragraph 0"]
        assert _paragraph_texts(path) == ["World"]

    def test_out_of_range_index_returns_error(self, tmp_path: Path) -> None:
        path = _make_test_doc(tmp_path / "bounds.docx")

        result = json.loads(
            write_word(path.name, operations=[{"action": "replace", "index": 9, "text": "Never"}])
        )

        assert result["applied_count"] == 0
        assert "超出范围" in result["errors"][0]
        assert _paragraph_texts(path) == ["Hello", "World"]

    def test_multiple_operations_batch(self, tmp_path: Path) -> None:
        path = _make_test_doc(tmp_path / "batch.docx", ["A", "B"])

        result = json.loads(
            write_word(
                path.name,
                operations=[
                    {"action": "replace", "index": 0, "text": "A1"},
                    {"action": "insert_after", "index": 0, "text": "A2"},
                    {"action": "delete", "index": 2},
                    {"action": "append", "text": "C"},
                ],
            )
        )

        assert result["applied_count"] == 4
        assert _paragraph_texts(path) == ["A1", "A2", "C"]


class TestInspectWord:
    def test_single_file(self, tmp_path: Path) -> None:
        path = tmp_path / "single.docx"
        doc = Document()
        doc.core_properties.title = "Quarterly Review"
        doc.add_heading("Overview", level=1)
        doc.add_paragraph("Body")
        doc.add_table(rows=1, cols=2)
        doc.save(path)

        result = json.loads(inspect_word(file_path=path.name))

        assert result["file"] == path.name
        assert result["title"] == "Quarterly Review"
        assert result["total_paragraphs"] == 2
        assert result["total_tables"] == 1
        assert result["headings"][0]["text"] == "Overview"

    def test_multiple_files(self, tmp_path: Path) -> None:
        first = _make_test_doc(tmp_path / "first.docx", ["One"])
        second = _make_test_doc(tmp_path / "second.docx", ["Two"])

        result = json.loads(inspect_word(file_paths=[first.name, second.name]))

        files = {entry["file"] for entry in result["files"]}
        assert files == {first.name, second.name}

    def test_directory_scan(self, tmp_path: Path) -> None:
        _make_test_doc(tmp_path / "a.docx")
        _make_test_doc(tmp_path / "b.docx")
        (tmp_path / "notes.txt").write_text("skip me", encoding="utf-8")

        result = json.loads(inspect_word(directory="."))

        files = {entry["file"] for entry in result["files"]}
        assert files == {"a.docx", "b.docx"}

    def test_directory_scan_skips_doc_files(self, tmp_path: Path) -> None:
        _make_test_doc(tmp_path / "current.docx", ["Current"])
        (tmp_path / "legacy.doc").write_bytes(b"legacy-doc")

        result = json.loads(inspect_word(directory="."))

        assert result["file"] == "current.docx"

    def test_non_docx_file_returns_error(self, tmp_path: Path) -> None:
        (tmp_path / "notes.txt").write_text("plain text", encoding="utf-8")

        result = json.loads(inspect_word(file_paths=["notes.txt"]))

        assert result["file"] == "notes.txt"
        assert "docx" in result["error"].lower()


class TestSearchWord:
    def test_contains(self, tmp_path: Path) -> None:
        path = _make_test_doc(tmp_path / "contains.docx", ["Hello world", "Bye"])

        result = json.loads(search_word("world", file_path=path.name))

        assert result["match_mode"] == "contains"
        assert result["total_matches"] == 1
        assert result["matches"][0]["paragraph_index"] == 0

    def test_exact(self, tmp_path: Path) -> None:
        path = _make_test_doc(tmp_path / "exact.docx", ["world", "worldwide"])

        result = json.loads(search_word("world", file_path=path.name, match_mode="exact"))

        assert result["total_matches"] == 1
        assert result["matches"][0]["text"] == "world"

    def test_regex(self, tmp_path: Path) -> None:
        path = _make_test_doc(tmp_path / "regex.docx", ["Item 42", "Other"])

        result = json.loads(search_word(r"Item \d+", file_path=path.name, match_mode="regex"))

        assert result["total_matches"] == 1
        assert result["matches"][0]["text"] == "Item 42"

    def test_startswith(self, tmp_path: Path) -> None:
        path = _make_test_doc(tmp_path / "startswith.docx", ["Intro line", "Tail"])

        result = json.loads(search_word("Intro", file_path=path.name, match_mode="startswith"))

        assert result["total_matches"] == 1
        assert result["matches"][0]["text"] == "Intro line"

    def test_multiple_files(self, tmp_path: Path) -> None:
        first = _make_test_doc(tmp_path / "first.docx", ["alpha"])
        second = _make_test_doc(tmp_path / "second.docx", ["beta alpha"])

        result = json.loads(search_word("alpha", file_paths=[first.name, second.name]))

        assert result["total_matches"] == 2
        files = {match["file"] for match in result["matches"]}
        assert files == {first.name, second.name}

    def test_doc_file_returns_error(self, tmp_path: Path) -> None:
        (tmp_path / "legacy.doc").write_bytes(b"legacy-doc")

        result = json.loads(search_word("legacy", file_path="legacy.doc"))

        assert result["file_path"] == "legacy.doc"
        assert "docx" in result["error"].lower()


class TestFileRegistryWordPolicy:
    def test_scan_workspace_does_not_classify_doc_as_word(self, tmp_path: Path) -> None:
        registry = FileRegistry(Database(str(tmp_path / "registry.db")), tmp_path)
        _make_test_doc(tmp_path / "report.docx", ["Intro"])
        (tmp_path / "legacy.doc").write_bytes(b"legacy-doc")

        registry.scan_workspace()
        report = registry.get_by_path("report.docx")
        legacy = registry.get_by_path("legacy.doc")
        assert report is not None
        assert report.file_type == "word"
        assert legacy is not None
        assert legacy.file_type == "other"


class TestHelpers:
    def test_ensure_docx_validation(self) -> None:
        assert _ensure_docx("report.docx") is None
        error = json.loads(_ensure_docx("report.txt"))
        assert error["file_path"] == "report.txt"

    def test_heading_level_parsing(self) -> None:
        doc = Document()
        title = doc.add_paragraph("Main Title")
        title.style = "Title"
        heading = doc.add_paragraph("Section")
        heading.style = "Heading 2"

        assert _heading_level(title) == 0
        assert _heading_level(heading) == 2

    def test_run_to_dict_format(self) -> None:
        para = Document().add_paragraph()
        run = para.add_run("Styled")
        run.bold = True
        run.italic = True
        run.underline = True
        run.font.name = "Calibri"
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)

        result = _run_to_dict(run)

        assert result == {
            "text": "Styled",
            "format": {
                "bold": True,
                "italic": True,
                "underline": True,
                "size_pt": 14.0,
                "font": "Calibri",
                "color": "AABBCC",
            },
        }
